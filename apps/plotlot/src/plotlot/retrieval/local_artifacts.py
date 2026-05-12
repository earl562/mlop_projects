"""Local document/spreadsheet artifact generation.

This replaces the prior hosted office-suite integration with deterministic,
server-local files built from open Office formats:

- ``python-docx`` creates ``.docx`` documents.
- ``openpyxl`` creates ``.xlsx`` workbooks.

The functions intentionally keep the chat-tool surface small: callers receive
a filename plus an API download URL instead of a third-party share URL.
"""

from __future__ import annotations

import re
import uuid
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter

from plotlot.config import settings


@dataclass(frozen=True)
class LocalArtifactResult:
    """Generated local artifact metadata."""

    filename: str
    file_path: str
    artifact_url: str
    title: str
    content_type: str
    size_bytes: int


def artifact_root() -> Path:
    """Return the configured artifact directory, creating it if needed."""

    root = Path(settings.artifact_storage_dir).expanduser()
    if not root.is_absolute():
        root = Path.cwd() / root
    root.mkdir(parents=True, exist_ok=True)
    return root


def artifact_path(filename: str) -> Path:
    """Resolve a safe artifact filename under the artifact root."""

    safe_name = Path(filename).name
    if not safe_name or safe_name in {".", ".."}:
        raise ValueError("Invalid artifact filename")
    return artifact_root() / safe_name


def _slugify(value: str) -> str:
    slug = re.sub(r"[^a-zA-Z0-9]+", "-", value.strip()).strip("-").lower()
    return slug[:56] or "plotlot-artifact"


def _artifact_metadata(path: Path, title: str, content_type: str) -> LocalArtifactResult:
    filename = path.name
    return LocalArtifactResult(
        filename=filename,
        file_path=str(path),
        artifact_url=f"/api/v1/artifacts/{filename}",
        title=title,
        content_type=content_type,
        size_bytes=path.stat().st_size,
    )


async def create_spreadsheet(
    title: str,
    headers: list[str],
    rows: list[list[str]],
) -> LocalArtifactResult:
    """Create a local ``.xlsx`` workbook with structured tabular data."""

    wb = Workbook()
    ws = wb.active
    ws.title = "Export"

    header_fill = PatternFill(start_color="B45309", end_color="B45309", fill_type="solid")
    header_font = Font(bold=True, color="FFFFFF")
    thin_border = Border(
        left=Side(style="thin", color="E7E5E4"),
        right=Side(style="thin", color="E7E5E4"),
        top=Side(style="thin", color="E7E5E4"),
        bottom=Side(style="thin", color="E7E5E4"),
    )

    for col_idx, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col_idx, value=header)
        cell.fill = header_fill
        cell.font = header_font
        cell.border = thin_border
        cell.alignment = Alignment(horizontal="center")

    for row_idx, row in enumerate(rows, 2):
        for col_idx, value in enumerate(row, 1):
            cell = ws.cell(row=row_idx, column=col_idx, value=value)
            cell.border = thin_border

    ws.freeze_panes = "A2"
    for col_idx, header in enumerate(headers, 1):
        values = [str(header)] + [
            str(row[col_idx - 1]) for row in rows if col_idx - 1 < len(row)
        ]
        width = min(max(len(value) for value in values) + 2, 48)
        ws.column_dimensions[get_column_letter(col_idx)].width = width

    filename = f"{_slugify(title)}-{uuid.uuid4().hex[:8]}.xlsx"
    path = artifact_root() / filename
    wb.save(path)
    return _artifact_metadata(
        path,
        title,
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )


async def create_document(title: str, content: str) -> LocalArtifactResult:
    """Create a local ``.docx`` document from plain text content."""

    doc = Document()
    doc.add_heading(title, level=1)

    blocks = [block.strip() for block in content.split("\n\n") if block.strip()]
    if not blocks and content.strip():
        blocks = [content.strip()]
    for block in blocks:
        for line in block.splitlines():
            text = line.strip()
            if text:
                doc.add_paragraph(text)
        doc.add_paragraph("")

    filename = f"{_slugify(title)}-{uuid.uuid4().hex[:8]}.docx"
    path = artifact_root() / filename
    doc.save(str(path))
    return _artifact_metadata(
        path,
        title,
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
